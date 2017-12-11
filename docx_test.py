import docx
from openpyxl import Workbook
import re
from openpyxl import load_workbook

# wb = Workbook()
# ws = wb.create_sheet(title="test")
# file = docx.Document("D:\\answer\\one.docx")
# print("段落数:"+str(len(file.paragraphs)))
#
# que_list=[]
# one = []
# pattern = re.compile(r'^.*\uff08([A-Z]+)\uff09.*$')
#
# for i in range(len(file.paragraphs)):
#     if len(one)==2:
#         txt = file.paragraphs[i-2].text.replace('\u3000', '').replace('\t', '').replace(' ', '')
#         match = pattern.match(txt)
#         if match:
#             one.append(match.group(1))
#         que_list.append(one)
#         one = []
#     txt = file.paragraphs[i].text.replace('\u3000', '').replace('\t', '').replace(' ', '')
#     one.append(txt)
# for que in que_list:
#     # print(que)
#     ws.append(que)
#     wb.save("D:\\answer\\one.xlsx")

# wb = load_workbook(filename='D:\\answer\\one.xlsx')
# test = wb.get_sheet_by_name('test')
# for row in range(1,test.max_row+1):
#     # print(test.cell(row = row,column=1).value.split('.')[1])
#     sp_list =test.cell(row = row,column=1).value.split('.')
#     if(len(sp_list)>1):
#         test.cell(row=row, column=1).value=sp_list[1]
#         print(test.cell(row=row, column=1).value)
#     # print(test.cell(row=row, column=2).value)
# wb.save('D:\\answer\\one.xlsx')